from abc import ABC, abstractmethod
from docx import Document
from io import BytesIO
import PyPDF2

class BaseDocumentProcessor(ABC):
    @abstractmethod
    def extract_text(self, file: BytesIO) -> str:
        pass

class PDFProcessor(BaseDocumentProcessor):
    def extract_text(self, file: BytesIO) -> str:
        reader = PyPDF2.PdfReader(file)
        text = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text.append(page_text)
                
        full_text = "\n".join(text)
        full_text = full_text.replace("-\n", "")
        import re
        full_text = re.sub(r'\s+', ' ', full_text)
        full_text = re.sub(r'\n\s*\n', '\n\n', full_text)
        
        return full_text

class DOCXProcessor(BaseDocumentProcessor):
    def extract_text(self, file: BytesIO) -> str:
        doc = Document(file)
        text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n".join(text)
    
class TXTProcessor(BaseDocumentProcessor):
    def extract_text(self, file: BytesIO) -> str:
        content = file.read()
        text = content.decode('utf-8', errors='replace') if isinstance(content, bytes) else content
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text

class DocumentProcessorFactory:
    processors = {
        'application/pdf': PDFProcessor,
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': DOCXProcessor,
        'text/plain': TXTProcessor 
    }

    @classmethod
    def get_processor(cls, mime_type: str) -> BaseDocumentProcessor:
        processor = cls.processors.get(mime_type)
        if not processor:
            raise ValueError(f"Unsupported file type: {mime_type}")
        return processor()